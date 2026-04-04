#!/usr/bin/env python3
"""Verifier for regulatory_submission_format task."""

import sys
import os
import logging

sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), '../../', 'utils'))
from wps_verification_utils import (
    copy_and_parse_document,
    cleanup_verification_temp,
    get_document_text,
    check_text_formatting,
    check_heading_styles,
    count_headings_by_level,
    count_tables,
    get_table_dimensions,
    vlm_verify_screenshot,
)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def verify_regulatory_submission_format(traj, env_info, task_info):
    """
    Verify that the PBRER document was formatted according to ICH E2C(R2) standards.

    PREREQUISITE: Core PBRER content must be preserved.

    SCORING CRITERIA:
    1. Section headings: ICH-mandated sections have heading styles applied
    2. Heading hierarchy: At least 2 heading levels used properly
    3. Tables created: Adverse event data converted from prose to tables
    4. Font specification: Body text uses Times New Roman
    5. Font size: Body text is 12pt
    6. Line spacing: Body text has double spacing
    7. Document control: Report number PBRER-2024-NXV-003 visible
    8. Content completeness: All major sections preserved
    9. VLM visual verification
    """
    copy_from_env = env_info.get('copy_from_env')
    if not copy_from_env:
        return {"passed": False, "score": 0, "feedback": "Copy function not available"}

    container_path = "/home/ga/Documents/pbrer_draft.docx"
    success, doc, error, temp_dir = copy_and_parse_document(
        container_path, copy_from_env, file_format='docx'
    )

    if not success:
        return {"passed": False, "score": 0, "feedback": error}

    try:
        feedback_parts = []
        full_text = get_document_text(doc).lower()

        # PREREQUISITE: PBRER content must be preserved
        key_phrases = [
            "nexovant",
            "casirivimab",
            "benefit-risk",
            "hepatotoxicity",
            "reporting interval",
            "marketing authorization",
        ]
        preserved = sum(1 for p in key_phrases if p.lower() in full_text)
        if preserved < 4:
            return {
                "passed": False,
                "score": 0,
                "feedback": f"PREREQUISITE FAILED: Content corrupted ({preserved}/{len(key_phrases)} key phrases)",
            }
        feedback_parts.append(f"Prerequisite: content preserved ({preserved}/{len(key_phrases)})")

        criteria_passed = 0
        total_criteria = 9

        # Criterion 1: ICH-mandated section headings applied
        ich_sections = {
            "executive summary": "Heading",
            "introduction": "Heading",
            "worldwide marketing authorization": "Heading",
            "patient exposure": "Heading",
            "signal evaluation": "Heading",
            "benefit-risk analysis": "Heading",
            "conclusion": "Heading",
        }
        sections_with_headings = 0
        for section_text, _ in ich_sections.items():
            for para in doc.paragraphs:
                if section_text in para.text.lower():
                    if para.style and 'heading' in para.style.name.lower():
                        sections_with_headings += 1
                    break

        if sections_with_headings >= 5:
            criteria_passed += 1
            feedback_parts.append(f"Section headings: {sections_with_headings}/{len(ich_sections)} applied")
        else:
            feedback_parts.append(f"Section headings: only {sections_with_headings}/{len(ich_sections)} (need 5+)")

        # Criterion 2: Heading hierarchy - at least 2 levels used
        heading_counts = count_headings_by_level(doc)
        levels_used = len(heading_counts)
        total_headings = sum(heading_counts.values())

        if levels_used >= 2 and total_headings >= 6:
            criteria_passed += 1
            feedback_parts.append(f"Heading hierarchy: {levels_used} levels, {total_headings} total headings")
        elif total_headings >= 6:
            feedback_parts.append(f"Heading hierarchy: {total_headings} headings but only {levels_used} level(s)")
        else:
            feedback_parts.append(f"Heading hierarchy: insufficient ({total_headings} headings, {levels_used} levels)")

        # Criterion 3: Tables created from adverse event prose data
        table_count = count_tables(doc)
        if table_count >= 2:
            criteria_passed += 1
            feedback_parts.append(f"Tables: {table_count} tables created (2+ required)")
        elif table_count == 1:
            feedback_parts.append(f"Tables: only {table_count} table (need 2+)")
        else:
            feedback_parts.append("Tables: NO tables created from prose data")

        # Criterion 4: Font specification - Times New Roman
        times_count = 0
        total_checked = 0
        for para in doc.paragraphs:
            if len(para.text.strip()) > 50:
                total_checked += 1
                for run in para.runs:
                    if run.font.name and 'times' in run.font.name.lower():
                        times_count += 1
                        break

        if total_checked > 0 and times_count >= total_checked * 0.6:
            criteria_passed += 1
            feedback_parts.append(f"Font: Times New Roman ({times_count}/{total_checked} paragraphs)")
        else:
            feedback_parts.append(f"Font: only {times_count}/{total_checked} paragraphs use Times New Roman")

        # Criterion 5: Font size - 12pt
        size_12_count = 0
        size_total = 0
        for para in doc.paragraphs:
            if len(para.text.strip()) > 50:
                for run in para.runs:
                    if run.text.strip():
                        size_total += 1
                        if run.font.size and abs(run.font.size.pt - 12.0) <= 1.0:
                            size_12_count += 1
                        elif run.font.size is None:
                            # Default font size is typically 11pt in docx
                            pass
                        break

        if size_total > 0 and size_12_count >= size_total * 0.5:
            criteria_passed += 1
            feedback_parts.append(f"Font size: 12pt ({size_12_count}/{size_total} runs)")
        else:
            feedback_parts.append(f"Font size: only {size_12_count}/{size_total} runs are 12pt")

        # Criterion 6: Double line spacing
        from docx.enum.text import WD_LINE_SPACING
        double_spaced = 0
        spacing_checked = 0
        for para in doc.paragraphs:
            if len(para.text.strip()) > 50:
                spacing_checked += 1
                pf = para.paragraph_format
                is_double = False
                if pf.line_spacing_rule == WD_LINE_SPACING.DOUBLE:
                    is_double = True
                elif pf.line_spacing is not None and isinstance(pf.line_spacing, float) and 1.9 <= pf.line_spacing <= 2.1:
                    is_double = True
                elif pf.line_spacing is not None and isinstance(pf.line_spacing, (int, float)) and pf.line_spacing == 2.0:
                    is_double = True
                if is_double:
                    double_spaced += 1

        if spacing_checked > 0 and double_spaced >= spacing_checked * 0.5:
            criteria_passed += 1
            feedback_parts.append(f"Line spacing: double ({double_spaced}/{spacing_checked} paragraphs)")
        else:
            feedback_parts.append(f"Line spacing: only {double_spaced}/{spacing_checked} paragraphs double-spaced")

        # Criterion 7: Document control header with report number in bold or heading
        # Must be formatted as a document header, not just raw text
        report_num = "pbrer-2024-nxv-003"
        doc_control_formatted = False
        for para in doc.paragraphs:
            if report_num in para.text.lower():
                is_heading = para.style and 'heading' in para.style.name.lower()
                is_bold = any(run.bold for run in para.runs if run.text.strip())
                if is_heading or is_bold:
                    doc_control_formatted = True
                    break
        if doc_control_formatted:
            criteria_passed += 1
            feedback_parts.append("Document control: report number formatted as header")
        else:
            feedback_parts.append("Document control: report number not formatted as header (need bold or heading)")

        # Criterion 8: Table has formatted headers (bold or shading)
        # Checks that tables created from prose are professionally formatted
        table_formatted = False
        for t_idx in range(min(table_count, 5)):
            from wps_verification_utils import check_table_header_formatting
            header_fmt = check_table_header_formatting(doc, t_idx)
            if header_fmt['has_bold'] or header_fmt['has_shading']:
                table_formatted = True
                break
        if table_formatted:
            criteria_passed += 1
            feedback_parts.append("Table formatting: professional headers present")
        elif table_count > 0:
            feedback_parts.append("Table formatting: tables exist but no formatted headers")
        else:
            feedback_parts.append("Table formatting: no tables to format")

        # Criterion 9: VLM verification
        vlm_result = vlm_verify_screenshot(env_info, traj, """
Analyze this WPS Writer screenshot of a regulatory submission document. Answer in JSON:
{
    "has_section_headings": true/false,
    "has_tables": true/false,
    "appears_formally_formatted": true/false,
    "has_consistent_font": true/false,
    "has_document_header": true/false
}
Does the document show:
1. Clear section headings with hierarchy?
2. Data tables with structured information?
3. Formal document formatting (not raw text)?
4. Consistent font usage throughout?
5. A document header or title block?
""")
        if vlm_result is not None:
            has_headings = vlm_result.get("has_section_headings", False)
            formal = vlm_result.get("appears_formally_formatted", False)
            if has_headings and formal:
                criteria_passed += 1
                feedback_parts.append("VLM: formal regulatory formatting confirmed")
            else:
                feedback_parts.append("VLM: formal formatting not confirmed")
        else:
            total_criteria -= 1
            feedback_parts.append("VLM: unavailable (skipped)")

        score = int((criteria_passed / total_criteria) * 100)
        passed = score >= 55

        return {
            "passed": passed,
            "score": score,
            "feedback": " | ".join(feedback_parts),
        }

    except Exception as e:
        logger.error(f"Verification error: {e}", exc_info=True)
        return {"passed": False, "score": 0, "feedback": f"Verification error: {str(e)}"}
    finally:
        cleanup_verification_temp(temp_dir)
