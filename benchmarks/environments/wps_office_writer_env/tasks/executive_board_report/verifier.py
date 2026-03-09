#!/usr/bin/env python3
"""Verifier for executive_board_report task."""

import sys
import os
import logging

sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), '../../', 'utils'))
from wps_verification_utils import (
    copy_and_parse_document,
    cleanup_verification_temp,
    get_document_text,
    check_text_formatting,
    check_paragraph_alignment,
    count_tables,
    get_table_dimensions,
    get_table_content,
    check_heading_styles,
    count_headings_by_level,
    check_table_header_formatting,
    vlm_verify_screenshot,
)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def verify_executive_board_report(traj, env_info, task_info):
    """
    Verify that the quarterly board report was properly created.

    PREREQUISITE: Core departmental data must be preserved.

    SCORING CRITERIA:
    1. Title: Centered and bold with company name and quarter
    2. Executive summary: Synthesized summary near start of document
    3. Financial performance table: Revenue/COGS/margins with budget comparison
    4. KPI dashboard table: Multiple metrics with actuals vs targets
    5. Recommendations/action matrix table: Prioritized items with owners
    6. Heading hierarchy: Proper section structure
    7. Table formatting: Professional header formatting on tables
    8. Content completeness: Data from all 5 departments represented
    9. VLM visual verification
    """
    copy_from_env = env_info.get('copy_from_env')
    if not copy_from_env:
        return {"passed": False, "score": 0, "feedback": "Copy function not available"}

    container_path = "/home/ga/Documents/quarterly_raw.docx"
    success, doc, error, temp_dir = copy_and_parse_document(
        container_path, copy_from_env, file_format='docx'
    )

    if not success:
        return {"passed": False, "score": 0, "feedback": error}

    try:
        feedback_parts = []
        full_text = get_document_text(doc).lower()

        # PREREQUISITE: Departmental data must be preserved
        key_phrases = [
            "vertex dynamics",
            "48.2",              # Revenue
            "q3 2024",
            "oee",
            "book-to-bill",
        ]
        preserved = sum(1 for p in key_phrases if p.lower() in full_text)
        if preserved < 3:
            return {
                "passed": False,
                "score": 0,
                "feedback": f"PREREQUISITE FAILED: Content corrupted ({preserved}/{len(key_phrases)} key phrases)",
            }
        feedback_parts.append(f"Prerequisite: content preserved ({preserved}/{len(key_phrases)})")

        criteria_passed = 0
        total_criteria = 9

        # Criterion 1: Title - centered and bold with company/quarter info
        title_found = False
        for para in doc.paragraphs[:10]:  # Check first 10 paragraphs
            text_l = para.text.lower()
            has_company = 'vertex' in text_l
            has_quarter = 'q3' in text_l or 'quarter' in text_l or 'quarterly' in text_l
            has_report = 'report' in text_l or 'board' in text_l

            if has_company and (has_quarter or has_report):
                is_bold = any(
                    run.bold for run in para.runs if run.text.strip()
                ) if para.runs else False
                is_centered = False
                if para.alignment is not None:
                    from docx.enum.text import WD_ALIGN_PARAGRAPH as WD
                    is_centered = para.alignment == WD.CENTER
                # Also check if it's a heading style (inherently formatted)
                is_heading = para.style and 'heading' in para.style.name.lower()

                if (is_bold or is_heading) and is_centered:
                    title_found = True
                    criteria_passed += 1
                    feedback_parts.append("Title: centered and bold with company/quarter")
                elif is_bold or is_centered or is_heading:
                    title_found = True
                    feedback_parts.append("Title: present but not fully formatted")
                break

        if not title_found:
            feedback_parts.append("Title: NOT found in first 10 paragraphs")

        # Criterion 2: Executive summary synthesizing findings
        # Must be a heading-styled paragraph or contain "executive summary" specifically
        # (not just "summary" which matches email subject lines in raw text)
        exec_summary = False
        for i, para in enumerate(doc.paragraphs):
            text_l = para.text.lower()
            is_heading = para.style and 'heading' in para.style.name.lower()
            has_exec_summary = 'executive summary' in text_l
            has_overview_heading = ('overview' in text_l or 'highlights' in text_l) and is_heading
            if has_exec_summary or has_overview_heading:
                if i < len(doc.paragraphs) * 0.3:
                    exec_summary = True
                    break

        if exec_summary:
            criteria_passed += 1
            feedback_parts.append("Executive summary: present near start")
        else:
            feedback_parts.append("Executive summary: NOT found near start of document")

        # Criterion 3: Financial performance table
        table_count = count_tables(doc)
        financial_table_found = False
        for t_idx in range(table_count):
            content = get_table_content(doc, t_idx)
            if not content:
                continue
            all_text = ' '.join(' '.join(row) for row in content).lower()
            has_financial = any(term in all_text for term in
                              ['revenue', 'cogs', 'cost of goods', 'gross margin',
                               'operating', 'net income', 'gross profit'])
            has_numbers = any(term in all_text for term in
                            ['48', '45', '31', '4.1', 'budget', 'actual', 'prior'])
            if has_financial and has_numbers:
                financial_table_found = True
                rows, cols = get_table_dimensions(doc, t_idx)
                feedback_parts.append(f"Financial table: found ({rows}x{cols})")
                break

        if financial_table_found:
            criteria_passed += 1
        else:
            feedback_parts.append("Financial table: NOT found (need revenue/costs with comparisons)")

        # Criterion 4: KPI dashboard table
        kpi_table_found = False
        for t_idx in range(table_count):
            content = get_table_content(doc, t_idx)
            if not content:
                continue
            all_text = ' '.join(' '.join(row) for row in content).lower()
            # KPI table should have metrics from multiple departments
            kpi_terms = ['oee', 'on-time delivery', 'scrap', 'turnover', 'satisfaction',
                        'safety', 'trir', 'book-to-bill', 'backlog', 'inventory',
                        'target', 'actual', 'kpi', 'metric']
            kpi_hits = sum(1 for t in kpi_terms if t in all_text)
            rows, cols = get_table_dimensions(doc, t_idx)
            if kpi_hits >= 3 and rows >= 5:
                kpi_table_found = True
                feedback_parts.append(f"KPI dashboard: found ({rows}x{cols}, {kpi_hits} KPI terms)")
                break

        if kpi_table_found:
            criteria_passed += 1
        else:
            feedback_parts.append("KPI dashboard: NOT found (need multi-department metrics table)")

        # Criterion 5: Recommendations/action matrix table
        action_table_found = False
        for t_idx in range(table_count):
            content = get_table_content(doc, t_idx)
            if not content:
                continue
            all_text = ' '.join(' '.join(row) for row in content).lower()
            has_action_terms = any(term in all_text for term in
                                 ['recommendation', 'action', 'priority', 'high', 'medium',
                                  'low', 'owner', 'timeline', 'initiative', 'strategic'])
            has_items = any(term in all_text for term in
                          ['building c', 'erp', 'itar', 'honeywell', 'succession',
                           'hvac', 'dod', 'expansion'])
            if has_action_terms and has_items:
                action_table_found = True
                rows, cols = get_table_dimensions(doc, t_idx)
                feedback_parts.append(f"Action matrix: found ({rows}x{cols})")
                break

        if action_table_found:
            criteria_passed += 1
        else:
            # Also accept if there's a recommendations section even without table
            if 'recommendation' in full_text or 'action item' in full_text:
                feedback_parts.append("Action matrix: recommendations text found but not in table form")
            else:
                feedback_parts.append("Action matrix: NOT found")

        # Criterion 6: Heading hierarchy
        heading_counts = count_headings_by_level(doc)
        levels_used = len(heading_counts)
        total_headings = sum(heading_counts.values())

        if levels_used >= 2 and total_headings >= 5:
            criteria_passed += 1
            feedback_parts.append(f"Heading hierarchy: {levels_used} levels, {total_headings} headings")
        elif total_headings >= 4:
            feedback_parts.append(f"Heading hierarchy: {total_headings} headings, {levels_used} level(s)")
        else:
            feedback_parts.append(f"Heading hierarchy: insufficient ({total_headings} headings)")

        # Criterion 7: Table formatting
        any_formatted = False
        for t_idx in range(min(table_count, 5)):
            header_fmt = check_table_header_formatting(doc, t_idx)
            if header_fmt['has_bold'] or header_fmt['has_shading']:
                any_formatted = True
                break

        if any_formatted:
            criteria_passed += 1
            feedback_parts.append("Table formatting: professional headers present")
        else:
            feedback_parts.append("Table formatting: no formatted table headers")

        # Criterion 8: Content completeness - all 5 departments represented
        departments = [
            ('finance', ['revenue', '48.2', 'net income', 'cash position', 'cogs']),
            ('operations', ['oee', 'production output', 'scrap rate', 'on-time delivery']),
            ('sales', ['bookings', 'book-to-bill', 'backlog', 'customer satisfaction']),
            ('engineering', ['r&d', 'patent', 'project phoenix', 'automated inspection']),
            ('hr', ['headcount', 'engagement', 'training', 'turnover', 'dei']),
        ]
        depts_represented = 0
        for dept_name, keywords in departments:
            if any(kw in full_text for kw in keywords):
                depts_represented += 1

        if depts_represented >= 4:
            criteria_passed += 1
            feedback_parts.append(f"Content completeness: {depts_represented}/5 departments")
        else:
            feedback_parts.append(f"Content completeness: only {depts_represented}/5 departments")

        # Criterion 9: VLM verification
        vlm_result = vlm_verify_screenshot(env_info, traj, """
Analyze this WPS Writer screenshot of a board report. Answer in JSON:
{
    "has_title": true/false,
    "has_tables": true/false,
    "has_section_headings": true/false,
    "appears_professionally_formatted": true/false
}
Does the document show:
1. A prominent title or header?
2. Data tables (financial, KPI, or action items)?
3. Clear section headings?
4. Professional board-level document formatting?
""")
        if vlm_result is not None:
            has_tables = vlm_result.get("has_tables", False)
            professional = vlm_result.get("appears_professionally_formatted", False)
            if has_tables or professional:
                criteria_passed += 1
                feedback_parts.append("VLM: board report format confirmed")
            else:
                feedback_parts.append("VLM: board report format not confirmed")
        else:
            total_criteria -= 1
            feedback_parts.append("VLM: unavailable (skipped)")

        score = int((criteria_passed / total_criteria) * 100)
        passed = score >= 55

        return {
            "passed": passed,
            "score": score,
            "feedback": " | ".join(feedback_parts),
        }

    except Exception as e:
        logger.error(f"Verification error: {e}", exc_info=True)
        return {"passed": False, "score": 0, "feedback": f"Verification error: {str(e)}"}
    finally:
        cleanup_verification_temp(temp_dir)
